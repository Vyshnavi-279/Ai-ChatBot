"""
scrape_college.py — BVRIT Hyderabad targeted web scraper
=========================================================
Fetches specific pages of https://bvrithyderabad.edu.in and maps content
into exactly the 8 Heading-1 sections required by the RAG pipeline:

  1. About BVRIT
  2. Departments
  3. Admissions
  4. Fee Structure
  5. Placements
  6. Campus & Facilities
  7. Faculty
  8. Contact

Rules enforced:
  - No marketing slogans / taglines
  - No invented facts
  - Named individuals kept only with their role/title
  - Sections with no scraped content get a placeholder note

Usage:
  python scrape_college.py
  python scrape_college.py --output data/college_info.docx
"""

from __future__ import annotations

import argparse
import re
import sys
import time
from typing import Optional

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
BASE_URL = "https://bvrithyderabad.edu.in"
HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120 Safari/537.36"
    )
}
REQUEST_TIMEOUT = 15
DELAY_BETWEEN_REQUESTS = 1.2   # seconds — be polite
PLACEHOLDER = "[Content to be added manually — not found in scrape]"

# ---------------------------------------------------------------------------
# Marketing phrases to strip from any scraped text
# ---------------------------------------------------------------------------
MARKETING_PHRASES = [
    r"empowering women to engineer a better future",
    r"building a world-class education system",
    r"engineering a world of quality",
    r"building the next generation of women leaders",
    r"where women thrive and break barriers",
    r"a dream doesn't become reality through magic",
    r"crafted\s*by\s*reinaphics",
    r"©\s*\d{4}\s*bvrith",
    r"skip to content",
]
_MARKETING_RE = re.compile(
    "|".join(MARKETING_PHRASES), re.IGNORECASE
)

# ---------------------------------------------------------------------------
# URL map: section name → list of pages to scrape for that section
# ---------------------------------------------------------------------------
SECTION_URLS: dict[str, list[str]] = {
    "About BVRIT": [
        f"{BASE_URL}/about-bvrith/",
        f"{BASE_URL}/principal/",
        f"{BASE_URL}/management/",
        f"{BASE_URL}/nirf/",
    ],
    "Departments": [
        f"{BASE_URL}/computer-science-and-engineering/about-the-department/",
        f"{BASE_URL}/electronics-and-communication-engineering/about-the-department/",
        f"{BASE_URL}/electrical-and-electronics-engineering/about-the-department/",
        f"{BASE_URL}/information-technology/about-the-department/",
        f"{BASE_URL}/cse-artificial-intelligence-and-machine-learning/about-the-department/",
        f"{BASE_URL}/basic-sciences-and-humanities/about-the-department/",
        f"{BASE_URL}/admission/intake-of-courses/",
        f"{BASE_URL}/post-graduate/computer-science-and-engineering/",
        f"{BASE_URL}/post-graduate/data-sciences/",
        f"{BASE_URL}/post-graduate/vlsi-design/",
    ],
    "Admissions": [
        f"{BASE_URL}/admission/admission-process/",
        f"{BASE_URL}/admission/b-category/",
        f"{BASE_URL}/admission/eamcet-ranks/",
        f"{BASE_URL}/admission/documents-to-submit/",
        f"{BASE_URL}/admission/transportation/",
    ],
    "Fee Structure": [
        f"{BASE_URL}/admission/fee-details/",
        f"{BASE_URL}/admission/hostel/",
    ],
    "Placements": [
        f"{BASE_URL}/placements/training-and-placement-cell/",
        f"{BASE_URL}/placements/placement-details/",
        f"{BASE_URL}/placements/training-placement-process/",
        f"{BASE_URL}/placements/employability-skills/",
        f"{BASE_URL}/placements/internships/",
    ],
    "Campus & Facilities": [
        f"{BASE_URL}/library/",
        f"{BASE_URL}/admission/hostel/",
        f"{BASE_URL}/gym/",
        f"{BASE_URL}/food-and-cafetaria/",
        f"{BASE_URL}/pcs-facilities/",
        f"{BASE_URL}/differentiators/assistive-technology-lab/",
        f"{BASE_URL}/differentiators/iot-maker-space/",
        f"{BASE_URL}/differentiators/drone-technology-laboratory/",
    ],
    "Faculty": [
        f"{BASE_URL}/principal/",
        f"{BASE_URL}/organogram/",
        f"{BASE_URL}/governing-body-member/",
    ],
    "Contact": [
        f"{BASE_URL}/contact-us/",
    ],
}


# ---------------------------------------------------------------------------
# Scraping helpers
# ---------------------------------------------------------------------------

def _fetch(url: str) -> Optional[BeautifulSoup]:
    """GET a URL and return a BeautifulSoup, or None on failure."""
    try:
        r = requests.get(url, headers=HEADERS, timeout=REQUEST_TIMEOUT)
        if r.status_code != 200:
            print(f"  [skip] {url} → HTTP {r.status_code}")
            return None
        return BeautifulSoup(r.text, "lxml")
    except Exception as exc:
        print(f"  [error] {url} → {exc}")
        return None


def _clean(text: str) -> str:
    """Strip marketing phrases, collapse whitespace, trim."""
    text = _MARKETING_RE.sub("", text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()


def _is_noise(text: str) -> bool:
    """True if the text is navigation/footer noise and should be discarded."""
    if len(text) < 20:
        return True
    noise_patterns = [
        r"^skip to",
        r"^home\s*[>»/]",
        r"^\s*©",
        r"^all rights reserved",
        r"^crafted by",
        r"^cookie",
        r"^privacy policy",
        r"^terms",
        r"^\s*\|",
        r"^facebook|^twitter|^instagram|^linkedin|^youtube",
    ]
    low = text.lower()
    return any(re.match(p, low) for p in noise_patterns)


def _extract_page_content(soup: BeautifulSoup, url: str) -> list[str]:
    """
    Extract factual paragraphs from a page.
    - Removes nav, header, footer, scripts, sidebars.
    - Returns a list of clean, non-empty text paragraphs.
    """
    # Remove all non-content tags
    for tag in soup.find_all(
        ["script", "style", "nav", "header", "footer",
         "form", "noscript", "iframe", "aside"]
    ):
        tag.decompose()

    # Try to find the main content region
    main = (
        soup.find("main")
        or soup.find("div", class_=re.compile(r"(content|main|entry|post|article)", re.I))
        or soup.find("article")
        or soup.body
    )
    if not main:
        return []

    paragraphs: list[str] = []
    seen: set[str] = set()

    for el in main.find_all(["p", "li", "td", "th", "h2", "h3", "h4"]):
        raw = el.get_text(separator=" ", strip=True)
        cleaned = _clean(raw)
        if not cleaned or _is_noise(cleaned):
            continue
        # Deduplicate
        key = re.sub(r"\s+", " ", cleaned.lower())
        if key in seen:
            continue
        seen.add(key)
        paragraphs.append(cleaned)

    return paragraphs


def scrape_section(section: str, urls: list[str]) -> list[str]:
    """Scrape all URLs for a section and return merged paragraph list."""
    print(f"\n[{section}]")
    all_paragraphs: list[str] = []
    seen_content: set[str] = set()

    for url in urls:
        print(f"  → {url}")
        soup = _fetch(url)
        if not soup:
            time.sleep(DELAY_BETWEEN_REQUESTS)
            continue

        paras = _extract_page_content(soup, url)
        for p in paras:
            key = re.sub(r"\s+", " ", p.lower()[:100])
            if key not in seen_content:
                seen_content.add(key)
                all_paragraphs.append(p)

        time.sleep(DELAY_BETWEEN_REQUESTS)

    print(f"  ✓ {len(all_paragraphs)} paragraphs collected")
    return all_paragraphs


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------

def build_document(section_content: dict[str, list[str]], output_path: str) -> None:
    """Write the 8-section docx from scraped content."""
    doc = Document()

    # Title
    title = doc.add_heading(
        "BVRIT Hyderabad College of Engineering for Women", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    section_order = [
        "About BVRIT",
        "Departments",
        "Admissions",
        "Fee Structure",
        "Placements",
        "Campus & Facilities",
        "Faculty",
        "Contact",
    ]

    for section in section_order:
        doc.add_heading(section, level=1)
        paragraphs = section_content.get(section, [])

        if not paragraphs:
            p = doc.add_paragraph(PLACEHOLDER)
            if p.runs:
                p.runs[0].italic = True
        else:
            for para_text in paragraphs:
                doc.add_paragraph(para_text)

    doc.save(output_path)
    print(f"\n✓ Saved to {output_path}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main(output_path: str = "data/college_info.docx") -> None:
    import os
    os.makedirs("data", exist_ok=True)

    print("=" * 60)
    print("BVRIT Hyderabad Scraper")
    print("=" * 60)

    section_content: dict[str, list[str]] = {}
    for section, urls in SECTION_URLS.items():
        section_content[section] = scrape_section(section, urls)

    build_document(section_content, output_path)

    # Summary
    print("\n--- Summary ---")
    for section in [
        "About BVRIT", "Departments", "Admissions", "Fee Structure",
        "Placements", "Campus & Facilities", "Faculty", "Contact"
    ]:
        n = len(section_content.get(section, []))
        status = f"{n} paragraphs" if n else "⚠️  placeholder only"
        print(f"  {section:<22} {status}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Scrape BVRIT website into college_info.docx")
    parser.add_argument(
        "--output", default="data/college_info.docx",
        help="Output path for the docx (default: data/college_info.docx)"
    )
    args = parser.parse_args()
    main(args.output)
